# -*- coding: utf-8 -*-
"""สร้างเอกสารหมวดที่ 4 หัวข้อ 4.3 คำอธิบายรายวิชา ตามแบบฟอร์มของมหาวิทยาลัย

ต้นทางทั้งหมดมาจากวอลต์ผ่านไฟล์ข้อมูลที่ sync แล้ว จึงสร้างใหม่ได้ทุกครั้งที่คำอธิบายรายวิชาเปลี่ยน
    python scripts/build-section4-3-docx.py            สร้างพร้อมคำอธิบายภาษาอังกฤษ
    python scripts/build-section4-3-docx.py --no-en    ตัดแถวคำอธิบายภาษาอังกฤษออก
"""
import json, re, sys, pathlib
from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = pathlib.Path(__file__).resolve().parents[2]
SRC = ROOT / "curriculum-graph" / "src"
OUT = ROOT / "curriculum-graph" / "data" / "exports" / "Section4_3_Course_Descriptions.docx"
WITH_EN = "--no-en" not in sys.argv

FONT, SIZE, SIZE_CLO = "TH Sarabun New", Pt(16), Pt(13.5)
COL_W = [Emu(977265), Emu(4215130), Emu(922655)]

# ─────────── อ่านข้อมูลจากไฟล์ที่ sync มาจากวอลต์ ───────────
def js_object(path, name):
    t = (SRC / path).read_text(encoding="utf-8")
    i = t.index("{", t.index("export const " + name)); depth = 0
    for j in range(i, len(t)):
        if t[j] == "{": depth += 1
        elif t[j] == "}":
            depth -= 1
            if depth == 0: break
    return json.loads(t[i:j + 1])

REVISION = js_object("courseRevisionData.js", "COURSE_REVISION")   # ชื่อ · หน่วยกิต · คำอธิบาย
KSEC = js_object("courseKsecData.js", "COURSE_KSEC")               # ระดับ Bloom ที่ระบุไว้ในวอลต์

# CLO ฉบับที่คณะกรรมการทวนสอบแล้ว (วิชาบังคับ) — เป็นชุดหลัก
VAULT_CLO = {}
vault_md = (ROOT / "Labor_Growth_Report_Vault" / "05_TQF2_Academic_Drafts" /
            "10_Course_Learning_Outcomes_CLO_Mapping.md").read_text(encoding="utf-8")
_cur = None
for line in vault_md.replace("\r\n", "\n").split("\n"):
    h = re.match(r"^###\s+(EN-714-\d{5})", line)
    if h:
        _cur = h.group(1); continue
    if _cur and re.match(r"^\|\s*CLO\d", line):
        cells = line.split("|")
        if len(cells) > 6:
            VAULT_CLO.setdefault(_cur, []).append(
                (re.sub(r"^CLO\d+\s*", "", cells[1].strip()), cells[4].strip()))

# CLO ของวิชาศึกษาทั่วไปและวิชาชีพเลือก — ยังไม่ผ่านการรับรอง ใช้เป็นชุดสำรอง
clo_js = (SRC / "cloData.js").read_text(encoding="utf-8")
CLO = {}
starts = [(m.start(), m.group(1)) for m in re.finditer(r'\{\s*c:\s*"([A-Z]{2}-\d{3}-\d{3,5})"', clo_js)]
for idx, (pos, code) in enumerate(starts):
    end = starts[idx + 1][0] if idx + 1 < len(starts) else len(clo_js)
    CLO[code] = [s.replace('\\"', '"').strip()
                 for s in re.findall(r'\bt:\s*"((?:[^"\\]|\\.)*)"', clo_js[pos:end])]

# ─────────── ระดับ Bloom ───────────
VERB = [("อธิบาย",2),("ยกตัวอย่าง",2),("สรุป",2),("บอก",2),("ตระหนัก",2),("ฟัง",2),("อ่าน",2),
        ("สืบค้น",3),("ระบุ",3),("ใช้",3),("เขียน",3),("คำนวณ",3),("จัดทำ",3),("ติดตั้ง",3),
        ("เชื่อมต่อ",3),("ต่อ",3),("ปฏิบัติ",3),("ดำเนินการ",3),("สื่อสาร",3),("นำเสนอ",3),
        ("จัดการ",3),("ควบคุม",3),("วางแผน",3),("นำ",3),("ประยุกต์",3),("สร้าง",3),("พัฒนา",3),
        ("ทำงาน",3),("เก็บ",3),("แสดง",3),("รายงาน",3),("สาธิต",3),("ทดลอง",3),("ตั้งค่า",3),
        ("ฝึก",3),("ปรับ",3),("ทำแผนที่",3),("วัด",3),("แบ่ง",3),
        ("ทดสอบ",3),("ประมวลผล",3),("แปลง",3),("บูรณาการ",3),("ติดตาม",3),("สอบเทียบ",3),
        ("ตีความ",2),("จำแนก",2),
        ("วิเคราะห์",4),("วินิจฉัย",4),("เลือก",4),("ตรวจสอบ",4),("พยากรณ์",4),
        ("กำหนด",4),("เปรียบเทียบ",4),("ออกแบบ",4),("ชี้บ่ง",4),("วิจัย",4),("ทวนสอบ",4),
        ("เสนอ",4),("วางกลยุทธ์",4),("บริหาร",4),
        ("ประเมิน",5),("ตัดสิน",5),("รับรอง",5),("สอบทาน",5),("นิยาม",5),("สังเคราะห์",6)]
UP5 = ("เกณฑ์การยอมรับ", "ตัดสิน", "เปรียบเทียบทางเลือก", "รับรอง")
UP6 = ("นิยามข้อกำหนดเอง", "โจทย์เปิด", "ยังไม่มีคำตอบ")
# ชื่อระดับแสดงเป็นภาษาไทยพร้อมภาษาอังกฤษ ไม่แสดงรหัส B ในเล่ม
NAME = {1: "จำ (Remember)", 2: "เข้าใจ (Understand)", 3: "ประยุกต์ใช้ (Apply)",
        4: "วิเคราะห์ (Analyze)", 5: "ประเมินค่า (Evaluate)", 6: "สร้างสรรค์ (Create)"}

def year_of(code):
    """ปีที่เปิดสอน — ใช้กำหนดเพดานระดับ Bloom"""
    m = re.match(r"EN-714-1(\d)(\d{3})", code)
    if not m: return None
    cat, seq = m.group(1), int(m.group(2))
    if cat == "1": return 1 if seq <= 9 else 2
    if cat == "2": return 1 if seq <= 2 else (2 if seq <= 10 else (3 if seq <= 18 else 4))
    if cat == "7": return 3 if seq == 1 else 4
    return None                                  # คลังวิชาเลือกไม่ผูกชั้นปี

def bloom(code, n, text):
    stated = next((c.get("bloom") for c in KSEC.get(code, {}).get("clos", []) if c.get("n") == n), None)
    if stated:
        return NAME.get(int(stated[1]), "—")
    hit = next((v for v in VERB if text.startswith(v[0])), None)
    if not hit: return "—"
    lvl = hit[1]
    if hit[0] in ("ออกแบบ", "พัฒนา", "สร้าง"):
        if any(k in text for k in UP6): lvl = 6
        elif any(k in text for k in UP5): lvl = 5
    y = year_of(code)
    ceiling = 3 if y == 1 else (4 if y == 2 else 5)
    return NAME[min(lvl, ceiling)]

# ─────────── การจัดกลุ่มตามหมวดรายวิชา ───────────
GROUPS = [
    ("1", lambda c: c[8] == "1", "กลุ่มวิชาแกนและพื้นฐานทางวิศวกรรม (หมวด 1 วิชาแกนหรือวิชาพื้นฐาน)"),
    ("2a", lambda c: c[8] == "2" and int(c[9:]) <= 8, "กลุ่มวิชาแกนปัญญาประดิษฐ์และระบบอัจฉริยะ (หมวด 2 วิชาชีพบังคับ)"),
    ("2b", lambda c: c[8] == "2" and 9 <= int(c[9:]) <= 16, "กลุ่มวิชาชีพบังคับตามโดเมน (หมวด 2 วิชาชีพบังคับ)"),
    ("2c", lambda c: c[8] == "2" and int(c[9:]) >= 17, "กลุ่มวิชาโครงงานและสัมมนา (หมวด 2 วิชาชีพบังคับ)"),
    ("4", lambda c: c[8] == "4", "กลุ่มวิชาชีพเลือก (หมวด 4 วิชาเลือก)"),
    ("6", lambda c: c[8] == "6", "กลุ่มวิชาฝึกปฏิบัติในสถานประกอบการ (หมวด 6)"),
    ("7", lambda c: c[8] == "7", "กลุ่มวิชาสหกิจศึกษา (หมวด 7)"),
]

# ─────────── สร้างเอกสาร ───────────
doc = Document()
st = doc.styles["Normal"]
st.font.name, st.font.size = FONT, SIZE
st.element.rPr.rFonts.set(__import__("docx").oxml.ns.qn("w:cs"), FONT)

def para(text, bold=False, size=SIZE, align=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align: p.alignment = align
    r = p.add_run(text); r.bold = bold
    r.font.name, r.font.size = FONT, size
    r.font.element.rPr.rFonts.set(__import__("docx").oxml.ns.qn("w:cs"), FONT)
    return p

def cell(c, text, bold=False, size=SIZE):
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text); r.bold = bold
    r.font.name, r.font.size = FONT, size
    r.font.element.rPr.rFonts.set(__import__("docx").oxml.ns.qn("w:cs"), FONT)

para("หมวดที่ 4 หัวข้อ 4.3 คำอธิบายรายวิชา", bold=True, size=Pt(18), align=WD_ALIGN_PARAGRAPH.CENTER)
para("หลักสูตรวิศวกรรมศาสตรบัณฑิต สาขาวิชาวิศวกรรมปัญญาประดิษฐ์และระบบอัจฉริยะ (หลักสูตรใหม่ พ.ศ. 2570)",
     align=WD_ALIGN_PARAGRAPH.CENTER)
para("สร้างจากคำอธิบายรายวิชาในคลังเอกสารหลักสูตรด้วยสคริปต์ build-section4-3-docx.py "
     "· ระดับ Bloom อ้างอิง Bloom's Revised Taxonomy โดยมิติความรู้ประเมินที่ B1–B2 และมิติทักษะที่ B3–B6",
     size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER)
para("หมายเหตุ ผลลัพธ์การเรียนรู้รายวิชาของกลุ่มวิชาชีพเลือกเป็นฉบับร่างเพื่อทวนสอบ "
     "ยังไม่ผ่านการรับรองจากคณะกรรมการบริหารหลักสูตรเช่นเดียวกับกลุ่มวิชาบังคับ",
     size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

total, no_clo = 0, []
for _, match, heading in GROUPS:
    codes = sorted(c for c in REVISION if match(c))
    if not codes: continue
    para(heading, bold=True, space_after=8)
    for code in codes:
        r = REVISION[code]
        vault = VAULT_CLO.get(code)
        clos = [x[0] for x in vault] if vault else CLO.get(code, [])
        stated = {i + 1: x[1] for i, x in enumerate(vault)} if vault else {}
        if not clos: no_clo.append(code)
        rows = 3 + (1 if WITH_EN else 0) + 1 + max(len(clos), 1)
        tb = doc.add_table(rows=rows, cols=3)
        tb.style = "Table Grid"
        for i, w in enumerate(COL_W):
            for row in tb.rows: row.cells[i].width = w

        cell(tb.rows[0].cells[0], code)
        cell(tb.rows[0].cells[1], r["t"])
        cell(tb.rows[0].cells[2], r["cr"])
        cell(tb.rows[1].cells[1], r["e"])
        cell(tb.rows[2].cells[1], r["d"])
        k = 3
        if WITH_EN:
            cell(tb.rows[k].cells[1], r.get("dEn", "")); k += 1
        head = tb.rows[k].cells
        head[0].merge(head[1]).paragraphs[0].text = ""
        cell(tb.rows[k].cells[0], "ผลลัพธ์การเรียนรู้รายวิชา (Course Learning Outcomes)", bold=True)
        k += 1
        for i, text in enumerate(clos, start=1):
            cell(tb.rows[k].cells[0], "CLO%d" % i, size=SIZE_CLO)
            cell(tb.rows[k].cells[1], text, size=SIZE_CLO)
            lv = stated.get(i)
            label = NAME[int(lv[1])] if lv else bloom(code, i, text)
            cell(tb.rows[k].cells[2], label, size=SIZE_CLO)
            k += 1
        if not clos:
            cell(tb.rows[k].cells[1], "— ยังไม่กำหนดผลลัพธ์การเรียนรู้รายวิชา —", size=SIZE_CLO)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        total += 1

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print("เขียน %s" % OUT.relative_to(ROOT))
print("  รายวิชา %d · มี CLO %d · ไม่มี CLO %d %s"
      % (total, total - len(no_clo), len(no_clo), no_clo[:5] if no_clo else ""))
print("  แถวคำอธิบายภาษาอังกฤษ: %s" % ("รวม" if WITH_EN else "ไม่รวม"))
